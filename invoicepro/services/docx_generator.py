"""
DOCX Invoice Generator
======================
Generates professional Word documents using python-docx.
Supports three visual presets: Classic, Modern, Minimal.
"""

from decimal import Decimal
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement


# ── Colour palettes per template ─────────────────────────────────────────────

PALETTES = {
    "classic": {
        "header_bg": RGBColor(0x1A, 0x23, 0x3B),   # dark navy
        "header_fg": RGBColor(0xFF, 0xFF, 0xFF),
        "accent": RGBColor(0x66, 0xE3, 0xFF),       # cyan
        "table_header_bg": RGBColor(0x1A, 0x23, 0x3B),
        "table_header_fg": RGBColor(0xFF, 0xFF, 0xFF),
        "table_alt": RGBColor(0xF2, 0xF6, 0xFE),
    },
    "modern": {
        "header_bg": RGBColor(0x6C, 0x63, 0xFF),    # violet
        "header_fg": RGBColor(0xFF, 0xFF, 0xFF),
        "accent": RGBColor(0x6C, 0x63, 0xFF),
        "table_header_bg": RGBColor(0x6C, 0x63, 0xFF),
        "table_header_fg": RGBColor(0xFF, 0xFF, 0xFF),
        "table_alt": RGBColor(0xF5, 0xF4, 0xFF),
    },
    "minimal": {
        "header_bg": RGBColor(0x22, 0x22, 0x22),    # near-black
        "header_fg": RGBColor(0xFF, 0xFF, 0xFF),
        "accent": RGBColor(0x22, 0x22, 0x22),
        "table_header_bg": RGBColor(0xEE, 0xEE, 0xEE),
        "table_header_fg": RGBColor(0x11, 0x11, 0x11),
        "table_alt": RGBColor(0xFA, 0xFA, 0xFA),
    },
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb.red:02X}{rgb.green:02X}{rgb.blue:02X}")
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, size=10, color: RGBColor | None = None,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _fmt(value) -> str:
    try:
        return f"{Decimal(value):,.2f}"
    except Exception:
        return str(value)


def _add_row(table, cols, alt: bool, palette):
    row = table.add_row()
    for i, (text, align, bold) in enumerate(cols):
        cell = row.cells[i]
        if alt:
            _set_cell_bg(cell, palette["table_alt"])
        _cell_text(cell, text, bold=bold, align=align)
    return row


# ── Main generator ────────────────────────────────────────────────────────────

def generate_docx(invoice) -> bytes:
    """
    Generate a DOCX document for the invoice.
    Returns the raw bytes of the .docx file.
    """
    template = getattr(invoice, "template", "classic") or "classic"
    if template not in PALETTES:
        template = "classic"
    pal = PALETTES[template]

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Header band ───────────────────────────────────────────────────────
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.style = "Table Grid"
    hdr_table.allow_autofit = True

    left_cell = hdr_table.cell(0, 0)
    right_cell = hdr_table.cell(0, 1)
    _set_cell_bg(left_cell, pal["header_bg"])
    _set_cell_bg(right_cell, pal["header_bg"])

    cs = invoice.user.company_settings
    company_name = cs.company_name or invoice.user.name
    _cell_text(left_cell, company_name, bold=True, size=16, color=pal["header_fg"])
    p2 = left_cell.add_paragraph()
    r2 = p2.add_run(cs.address or "")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)
    if cs.gstin:
        p3 = left_cell.add_paragraph()
        r3 = p3.add_run(f"GSTIN: {cs.gstin}")
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)

    _cell_text(right_cell, "INVOICE", bold=True, size=22, color=pal["accent"],
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    for label, value in [
        ("Invoice #", invoice.invoice_number),
        ("Date", invoice.invoice_date.strftime("%d %b %Y") if invoice.invoice_date else ""),
        ("Due", invoice.due_date.strftime("%d %b %Y") if invoice.due_date else ""),
        ("Status", invoice.status),
    ]:
        p = right_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f"{label}: {value}")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)

    doc.add_paragraph()

    # ── Bill To ───────────────────────────────────────────────────────────
    bt = doc.add_paragraph()
    bt.add_run("BILL TO").bold = True
    bt.runs[0].font.size = Pt(9)
    bt.runs[0].font.color.rgb = pal["accent"]

    cust = invoice.customer
    for line in [
        cust.name,
        cust.email or "",
        cust.phone or "",
        cust.address or "",
        f"GSTIN: {cust.gstin}" if cust.gstin else "",
    ]:
        if line:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)

    doc.add_paragraph()

    # ── Items table ───────────────────────────────────────────────────────
    col_widths = [Inches(2.4), Inches(0.7), Inches(0.9), Inches(0.9), Inches(0.8), Inches(0.9), Inches(0.9)]
    headers = ["Item / Description", "HSN/SAC", "Qty", "Unit Price", "Discount", "GST %", "Total"]
    alignments = [
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
    ]

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr_row = tbl.rows[0]
    for i, (hdr_text, align) in enumerate(zip(headers, alignments)):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, pal["table_header_bg"])
        _cell_text(cell, hdr_text, bold=True, size=9, color=pal["table_header_fg"], align=align)

    for idx, item in enumerate(invoice.items):
        alt = idx % 2 == 1
        cols = [
            (item.item_name + ("\n" + (item.description or "") if item.description else ""),
             WD_ALIGN_PARAGRAPH.LEFT, False),
            (item.hsn_sac or "", WD_ALIGN_PARAGRAPH.CENTER, False),
            (_fmt(item.quantity), WD_ALIGN_PARAGRAPH.CENTER, False),
            (_fmt(item.unit_price), WD_ALIGN_PARAGRAPH.RIGHT, False),
            (_fmt(item.discount), WD_ALIGN_PARAGRAPH.RIGHT, False),
            (f"{_fmt(item.gst_rate)}%", WD_ALIGN_PARAGRAPH.CENTER, False),
            (_fmt(item.line_total), WD_ALIGN_PARAGRAPH.RIGHT, True),
        ]
        _add_row(tbl, cols, alt, pal)

    doc.add_paragraph()

    # ── Totals block ──────────────────────────────────────────────────────
    totals_table = doc.add_table(rows=0, cols=2)
    totals_table.style = "Table Grid"

    def _totals_row(label, value, bold=False):
        row = totals_table.add_row()
        _cell_text(row.cells[0], label, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(row.cells[1], f"{invoice.currency} {value}", bold=bold,
                   align=WD_ALIGN_PARAGRAPH.RIGHT)

    _totals_row("Subtotal", _fmt(invoice.subtotal))
    _totals_row("Discount", f"- {_fmt(invoice.discount_total)}")
    _totals_row("Taxable Amount", _fmt(invoice.taxable_amount))
    if Decimal(invoice.cgst_amount or 0) > 0:
        _totals_row("CGST", _fmt(invoice.cgst_amount))
        _totals_row("SGST", _fmt(invoice.sgst_amount))
    if Decimal(invoice.igst_amount or 0) > 0:
        _totals_row("IGST", _fmt(invoice.igst_amount))
    _totals_row("TOTAL", _fmt(invoice.total_amount), bold=True)
    _totals_row("Amount Paid", _fmt(invoice.amount_paid))
    _totals_row("Balance Due", _fmt(invoice.balance_due), bold=True)

    doc.add_paragraph()

    # ── Notes / Payment ───────────────────────────────────────────────────
    if invoice.notes:
        doc.add_paragraph().add_run("Notes").bold = True
        doc.add_paragraph(invoice.notes)

    pi = cs.payment_instructions
    if pi:
        doc.add_paragraph().add_run("Payment Instructions").bold = True
        doc.add_paragraph(pi)

    if cs.footer_text:
        p = doc.add_paragraph(cs.footer_text)
        p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        p.runs[0].font.size = Pt(8)

    # ── Serialise ─────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
